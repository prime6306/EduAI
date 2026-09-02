"""
Generation -> validate/retry-once -> assemble multi-set paper -> save ->
auto-save every question to the question bank (with embedding-similarity
dedup) -> MLflow log.
"""
import random
import time
from datetime import datetime

from bson import ObjectId
from bson.errors import InvalidId
from flask import current_app

from app.extensions import db, logger
from . import prompts
from app.modules.nlp.llm_client import chat_json

DUPLICATE_SIMILARITY_THRESHOLD = 0.92


def _normalize_questions(raw):
    out = {"mcq": [], "short": [], "long": []}
    for key in ("mcq", "short", "long"):
        for q in raw.get(key, []):
            if not q.get("text"):
                continue
            if key == "mcq":
                options = q.get("options", [])
                if len(options) != 4:
                    continue
                correct_index = q.get("correct_index")
                if not isinstance(correct_index, int) or not (0 <= correct_index < 4):
                    continue
            out[key].append(q)
    if not any(out.values()):
        raise ValueError("Generated paper has no valid questions.")
    return out


def _generate_with_retry(messages):
    try:
        raw = chat_json(messages)
        return _normalize_questions(raw)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Question paper generation malformed, retrying once: %s", exc)
        raw = chat_json(messages)
        return _normalize_questions(raw)


def _shuffle_mcq_options(q):
    options = list(q["options"])
    correct_text = options[q["correct_index"]]
    shuffled = options[:]
    random.shuffle(shuffled)
    new_q = dict(q)
    new_q["options"] = shuffled
    new_q["correct_index"] = shuffled.index(correct_text)
    return new_q


def _assign_marks_and_type(questions, q_type, marks):
    out = []
    for q in questions:
        item = dict(q)
        item["type"] = q_type
        item["marks"] = marks
        out.append(item)
    return out


def _build_set(sections, set_label, shuffle):
    questions = []
    for q_type in ("mcq", "short", "long"):
        section_qs = list(sections[q_type])
        if shuffle:
            random.shuffle(section_qs)
            section_qs = [_shuffle_mcq_options(q) if q_type == "mcq" else q for q in section_qs]
        questions.extend(section_qs)
    return {"set_label": set_label, "questions": questions}


def generate_paper(
    teacher_id, subject, units, total_marks,
    difficulty_split, question_types, marks_per_type,
    section_marks, num_sets, bloom_levels, duration,
):
    start = time.time()

    mcq_count = section_marks.get("mcq", 0) // marks_per_type.get("mcq", 1) if "mcq" in question_types else 0
    short_count = section_marks.get("short", 0) // marks_per_type.get("short", 1) if "short" in question_types else 0
    long_count = section_marks.get("long", 0) // marks_per_type.get("long", 1) if "long" in question_types else 0

    messages = prompts.question_paper_prompt(
        subject, units, mcq_count, short_count, long_count,
        marks_per_type.get("mcq", 1), marks_per_type.get("short", 5), marks_per_type.get("long", 15),
        difficulty_split, bloom_levels,
    )
    normalized = _generate_with_retry(messages)

    sections = {
        "mcq": _assign_marks_and_type(normalized["mcq"][:mcq_count], "mcq", marks_per_type.get("mcq", 1)),
        "short": _assign_marks_and_type(normalized["short"][:short_count], "short", marks_per_type.get("short", 5)),
        "long": _assign_marks_and_type(normalized["long"][:long_count], "long", marks_per_type.get("long", 15)),
    }

    set_labels = ["Set A", "Set B", "Set C"][:max(1, min(num_sets, 3))]
    sets = [_build_set(sections, label, shuffle=(i > 0)) for i, label in enumerate(set_labels)]

    actual_marks = sum(q["marks"] for q in sets[0]["questions"])

    doc = {
        "teacher_id": teacher_id,
        "subject": subject,
        "units": units,
        "total_marks": actual_marks,
        "requested_total_marks": total_marks,
        "difficulty_split": difficulty_split,
        "duration": duration,
        "sets": sets,
        "created_at": datetime.utcnow(),
        "exported_at": None,
    }
    result = db.question_papers.insert_one(doc)
    doc["_id"] = result.inserted_id

    _save_to_question_bank(teacher_id, subject, sets[0]["questions"])
    _log_to_mlflow(subject, len(sets[0]["questions"]), difficulty_split, round(time.time() - start, 1))

    return doc


def _save_to_question_bank(teacher_id, subject, questions):
    try:
        from app.modules.rag.vector_store import embed_texts
        import numpy as np

        texts = [q["text"] for q in questions]
        embeddings = embed_texts(texts)

        existing = list(db.question_bank.find({"teacher_id": teacher_id, "subject": subject}, {"embedding": 1}))
        existing_vecs = np.array([e["embedding"] for e in existing]) if existing else None

        for q, emb in zip(questions, embeddings):
            if existing_vecs is not None and len(existing_vecs):
                sims = existing_vecs @ np.array(emb) / (
                    np.linalg.norm(existing_vecs, axis=1) * np.linalg.norm(emb) + 1e-9
                )
                if sims.max() > DUPLICATE_SIMILARITY_THRESHOLD:
                    continue

            db.question_bank.insert_one({
                "teacher_id": teacher_id, "subject": subject, "topic": q.get("topic", ""),
                "text": q["text"], "type": q["type"], "marks": q["marks"],
                "difficulty": q.get("difficulty", ""), "bloom_level": q.get("bloom_level", ""),
                "answer": q.get("answer") or (q["options"][q["correct_index"]] if q["type"] == "mcq" else ""),
                "options": q.get("options"), "correct_index": q.get("correct_index"),
                "embedding": emb, "created_at": datetime.utcnow(), "used_count": 1,
            })
    except Exception:  # noqa: BLE001
        logger.warning("Question bank save/dedup failed - continuing without it.", exc_info=True)


def _log_to_mlflow(subject, n_questions, difficulty_split, processing_time):
    try:
        import mlflow
        mlflow.set_tracking_uri(current_app.config["MLFLOW_TRACKING_URI"])
        mlflow.set_experiment(current_app.config["MLFLOW_EXPERIMENT"])
        with mlflow.start_run(run_name="question_paper_generation"):
            mlflow.log_param("subject", subject)
            mlflow.log_param("difficulty_split", str(difficulty_split))
            mlflow.log_metric("total_questions", n_questions)
            mlflow.log_metric("processing_time_sec", processing_time)
    except Exception:  # noqa: BLE001
        logger.info("MLflow not reachable - skipping question paper run log.")


def get_paper(paper_id, teacher_id):
    try:
        oid = ObjectId(paper_id)
    except (InvalidId, TypeError):
        return None
    return db.question_papers.find_one({"_id": oid, "teacher_id": teacher_id})


def list_papers(teacher_id):
    return list(db.question_papers.find({"teacher_id": teacher_id}).sort("created_at", -1))


def regenerate_question(paper_id, teacher_id, set_label, question_index):
    doc = get_paper(paper_id, teacher_id)
    if not doc:
        return None
    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), None)
    if not target_set or question_index >= len(target_set["questions"]):
        return None

    old_q = target_set["questions"][question_index]
    messages = prompts.single_question_prompt(
        doc["subject"], old_q.get("topic", doc["subject"]), old_q["type"], old_q["marks"],
        old_q.get("difficulty", "Medium"),
    )
    raw = chat_json(messages)
    new_q = dict(old_q)
    new_q["text"] = raw.get("text", old_q["text"])
    if old_q["type"] == "mcq":
        new_q["options"] = raw.get("options", old_q.get("options"))
        new_q["correct_index"] = raw.get("correct_index", old_q.get("correct_index"))
    else:
        new_q["answer"] = raw.get("answer", old_q.get("answer"))

    target_set["questions"][question_index] = new_q
    db.question_papers.update_one({"_id": doc["_id"]}, {"$set": {"sets": doc["sets"]}})
    return new_q


def edit_question(paper_id, teacher_id, set_label, question_index, fields):
    doc = get_paper(paper_id, teacher_id)
    if not doc:
        return False
    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), None)
    if not target_set or question_index >= len(target_set["questions"]):
        return False
    target_set["questions"][question_index].update(fields)
    db.question_papers.update_one({"_id": doc["_id"]}, {"$set": {"sets": doc["sets"]}})
    return True


def delete_question(paper_id, teacher_id, set_label, question_index):
    doc = get_paper(paper_id, teacher_id)
    if not doc:
        return False
    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), None)
    if not target_set or question_index >= len(target_set["questions"]):
        return False
    target_set["questions"].pop(question_index)
    new_total = sum(q["marks"] for q in target_set["questions"])
    db.question_papers.update_one(
        {"_id": doc["_id"]}, {"$set": {"sets": doc["sets"], "total_marks": new_total}}
    )
    return True


def add_manual_question(paper_id, teacher_id, set_label, question):
    doc = get_paper(paper_id, teacher_id)
    if not doc:
        return False
    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), None)
    if not target_set:
        return False
    target_set["questions"].append(question)
    new_total = sum(q["marks"] for q in target_set["questions"])
    db.question_papers.update_one(
        {"_id": doc["_id"]}, {"$set": {"sets": doc["sets"], "total_marks": new_total}}
    )
    return True


def list_question_bank(teacher_id, subject="", q_type="", difficulty=""):
    query = {"teacher_id": teacher_id}
    if subject:
        query["subject"] = {"$regex": subject, "$options": "i"}
    if q_type:
        query["type"] = q_type
    if difficulty:
        query["difficulty"] = difficulty
    return list(db.question_bank.find(query).sort("created_at", -1))


def delete_bank_question(question_id, teacher_id):
    try:
        oid = ObjectId(question_id)
    except (InvalidId, TypeError):
        return False
    result = db.question_bank.delete_one({"_id": oid, "teacher_id": teacher_id})
    return result.deleted_count > 0


def render_pdf(doc, set_label, show_answers):
    from weasyprint import HTML
    from flask import render_template

    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), doc["sets"][0])
    html_string = render_template(
        "question_paper/pdf.html", doc=doc, q_set=target_set, show_answers=show_answers,
    )
    return HTML(string=html_string).write_pdf()


def render_docx(doc, set_label, show_answers):
    import io
    from docx import Document

    target_set = next((s for s in doc["sets"] if s["set_label"] == set_label), doc["sets"][0])

    d = Document()
    d.add_heading(doc["subject"], level=1)
    d.add_paragraph(f"{target_set['set_label']} - Total Marks: {doc['total_marks']} - Duration: {doc.get('duration', '')}")

    for section_name, q_type in [("Section A - MCQ", "mcq"), ("Section B - Short Answer", "short"), ("Section C - Long Answer", "long")]:
        qs = [q for q in target_set["questions"] if q["type"] == q_type]
        if not qs:
            continue
        d.add_heading(section_name, level=2)
        for i, q in enumerate(qs, 1):
            d.add_paragraph(f"{i}. {q['text']} [{q['marks']} marks]")
            if q_type == "mcq":
                for j, opt in enumerate(q.get("options", [])):
                    marker = "*" if show_answers and j == q.get("correct_index") else " "
                    d.add_paragraph(f"   {marker}({chr(65+j)}) {opt}")
            elif show_answers and q.get("answer"):
                d.add_paragraph(f"   Answer: {q['answer']}")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
